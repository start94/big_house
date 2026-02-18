import uvicorn
import os
import sqlite3
import jwt
import json
from fastapi import FastAPI, Depends, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import FileResponse
from starlette.middleware.sessions import SessionMiddleware
from starlette.requests import Request
from starlette.responses import RedirectResponse
from starlette.config import Config
from authlib.integrations.starlette_client import OAuth
from pydantic import BaseModel, EmailStr
from datetime import datetime, date, timedelta, timezone
from enum import Enum
from passlib.context import CryptContext
from contextlib import contextmanager
from dotenv import load_dotenv
from docx import Document # Richiede: pip install python-docx
from docx.shared import Pt, RGBColor

# --- CONFIGURAZIONE ---
load_dotenv()

SECRET_KEY = os.getenv("SECRET_KEY", "chiave_super_segreta_per_demo_v2_lunga")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60
DATABASE_PATH = os.getenv("DATABASE_PATH", "bighouse.db")
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:5173")
BACKEND_URL = os.getenv("BACKEND_URL", "http://localhost:8000")

# Google Auth Config
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")

app = FastAPI(title="Big House API")

app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)

origins = [
    "http://localhost:5173",
    "http://localhost:3000",
    FRONTEND_URL
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="auth/token")

config_data = {'GOOGLE_CLIENT_ID': GOOGLE_CLIENT_ID, 'GOOGLE_CLIENT_SECRET': GOOGLE_CLIENT_SECRET}
starlette_config = Config(environ=config_data)
oauth = OAuth(starlette_config)

if GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET:
    oauth.register(
        name='google',
        server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
        client_kwargs={'scope': 'openid email profile'}
    )

class Plan(str, Enum):
    FREE = "free"
    PRO = "pro"
    PLUS = "plus"

class RenovationLevel(str, Enum):
    BASSO = "basso"
    MEDIO = "medio"
    ALTO = "alto"

# --- DATABASE SETUP ---
def init_db():
    conn = sqlite3.connect(DATABASE_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            name TEXT NOT NULL,
            hashed_password TEXT,
            plan TEXT DEFAULT 'free',
            usage_date TEXT,
            deepresearch_count INTEGER DEFAULT 0,
            calcola_count INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            auth_provider TEXT DEFAULT 'local'
        )
    """)
    try:
        cursor.execute("ALTER TABLE users ADD COLUMN auth_provider TEXT DEFAULT 'local'")
    except sqlite3.OperationalError:
        pass 
    conn.commit()
    conn.close()

@contextmanager
def get_db():
    conn = sqlite3.connect(DATABASE_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
    finally:
        conn.close()

init_db()

# --- MODELLI PYDANTIC ---
class UserRegister(BaseModel):
    email: EmailStr
    password: str
    name: str

class Token(BaseModel):
    access_token: str
    token_type: str

class UserOut(BaseModel):
    name: str
    email: EmailStr
    plan: Plan
    usage: dict

# MODIFICATO: Aggiunto livello ristrutturazione
class CalculationRequest(BaseModel):
    buy_price: float
    surface: float
    city: str
    renovation_level: RenovationLevel

class DeepResearchRequest(BaseModel):
    query: str

class PlanUpdate(BaseModel):
    plan: Plan

# --- LOGICA AUTH ---
def verify_password(plain_password, hashed_password):
    if not hashed_password: return False
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

# --- DATABASE HELPERS ---
def get_user_by_email(email: str):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE email = ?", (email,))
        row = cursor.fetchone()
        if row:
            return dict(row)
    return None

def create_user(email: str, name: str, hashed_password: str = None, provider: str = "local"):
    with get_db() as conn:
        cursor = conn.cursor()
        today = str(date.today())
        cursor.execute("""
            INSERT INTO users (email, name, hashed_password, plan, usage_date, deepresearch_count, calcola_count, auth_provider)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (email, name, hashed_password, "free", today, 0, 0, provider))
        conn.commit()
        return cursor.lastrowid

def update_user_plan(email: str, new_plan: str):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("UPDATE users SET plan = ? WHERE email = ?", (new_plan, email))
        conn.commit()

def reset_usage_if_new_day(user: dict):
    today = str(date.today())
    if user["usage_date"] != today:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE users 
                SET usage_date = ?, deepresearch_count = 0, calcola_count = 0
                WHERE email = ?
            """, (today, user["email"]))
            conn.commit()
        user["usage_date"] = today
        user["deepresearch_count"] = 0
        user["calcola_count"] = 0

def increment_usage(email: str, feature: str):
    with get_db() as conn:
        cursor = conn.cursor()
        if feature == "deepresearch":
            cursor.execute("UPDATE users SET deepresearch_count = deepresearch_count + 1 WHERE email = ?", (email,))
        else:
            cursor.execute("UPDATE users SET calcola_count = calcola_count + 1 WHERE email = ?", (email,))
        conn.commit()

async def get_current_user(token: str = Depends(oauth2_scheme)):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(status_code=401, detail="Credenziali non valide")
    except Exception:
        raise HTTPException(status_code=401, detail="Credenziali non valide")
    
    user = get_user_by_email(email)
    if user is None:
        raise HTTPException(status_code=401, detail="Utente non trovato")
    
    reset_usage_if_new_day(user)
    return user

def check_limit(user: dict, feature: str):
    plan = user["plan"]
    count = user[f"{feature}_count"]
    if plan == "free":
        raise HTTPException(status_code=403, detail="Upgrade richiesto")
    if plan == "pro" and count >= 2:
        raise HTTPException(status_code=429, detail="Limite giornaliero raggiunto")
    return True

# --- ENDPOINTS ---

@app.post("/auth/register", response_model=Token)
async def register(user_in: UserRegister):
    if get_user_by_email(user_in.email):
        raise HTTPException(status_code=400, detail="Email già registrata")
    
    hashed_pw = get_password_hash(user_in.password)
    create_user(user_in.email, user_in.name, hashed_pw, "local")
    
    access_token = create_access_token(data={"sub": user_in.email})
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/auth/token", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = get_user_by_email(form_data.username)
    if not user:
        raise HTTPException(status_code=400, detail="Utente non trovato")
    if user['auth_provider'] == 'google':
        raise HTTPException(status_code=400, detail="Usa il login con Google")
    if not verify_password(form_data.password, user["hashed_password"]):
        raise HTTPException(status_code=400, detail="Password errata")
    
    access_token = create_access_token(data={"sub": user["email"]})
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/auth/google")
async def login_google(request: Request):
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=500, detail="Google Auth non configurato")
    redirect_uri = f"{BACKEND_URL}/auth/google/callback"
    return await oauth.google.authorize_redirect(request, redirect_uri)

@app.get("/auth/google/callback")
async def auth_google(request: Request):
    try:
        token = await oauth.google.authorize_access_token(request)
    except Exception:
        raise HTTPException(status_code=400, detail="Errore Auth Google")
    
    user_info = token.get('userinfo')
    if not user_info:
        user_info = await oauth.google.userinfo(token=token)

    email = user_info.get("email")
    name = user_info.get("name")
    
    user = get_user_by_email(email)
    if not user:
        create_user(email, name, hashed_password=None, provider="google")
    
    access_token = create_access_token(data={"sub": email})
    return RedirectResponse(url=f"{FRONTEND_URL}?token={access_token}")

@app.get("/users/me", response_model=UserOut)
async def read_users_me(current_user: dict = Depends(get_current_user)):
    return {
        "name": current_user["name"],
        "email": current_user["email"],
        "plan": current_user["plan"],
        "usage": {
            "deepresearch": current_user["deepresearch_count"],
            "calcola": current_user["calcola_count"]
        }
    }

@app.post("/billing/upgrade")
async def upgrade_plan(plan_update: PlanUpdate, current_user: dict = Depends(get_current_user)):
    update_user_plan(current_user["email"], plan_update.plan.value)
    return {"status": "success", "new_plan": plan_update.plan}

# --- NUOVA LOGICA CALCOLO IMMOBILIARE ---
@app.post("/features/calculate")
async def calculate_roi(req: CalculationRequest, current_user: dict = Depends(get_current_user)):
    check_limit(current_user, "calcola")
    
    # 1. Parametri base per città (simulati per demo)
    city_multipliers = {
        "milano": 1.5, "roma": 1.3, "napoli": 1.0, "torino": 0.9, "firenze": 1.4, "parigi": 2.5
    }
    multiplier = city_multipliers.get(req.city.lower(), 1.0)
    
    # 2. Costi ristrutturazione al mq in base al livello
    renovation_costs = {
        "basso": 500,   # €/mq
        "medio": 800,   # €/mq
        "alto": 1200    # €/mq
    }
    cost_per_sqm = renovation_costs[req.renovation_level] * multiplier
    total_renovation_cost = cost_per_sqm * req.surface
    
    # 3. Calcolo Investimento Totale
    # Aggiungiamo un 10% di spese accessorie (notaio, agenzia, tasse)
    acquisition_costs = req.buy_price * 0.10
    total_investment = req.buy_price + total_renovation_cost + acquisition_costs
    
    # 4. Stima Valore Post-Ristrutturazione (Flip)
    # L'incremento di valore dipende dal livello di ristrutturazione
    value_increase_pct = {
        "basso": 0.20,  # +20%
        "medio": 0.35,  # +35%
        "alto": 0.55    # +55%
    }
    estimated_final_value = total_investment * (1 + value_increase_pct[req.renovation_level])
    
    # 5. Calcolo ROI e Tempi
    profit = estimated_final_value - total_investment
    roi_percentage = (profit / total_investment) * 100
    
    duration_months = {
        "basso": 3,
        "medio": 5,
        "alto": 8
    }[req.renovation_level]

    if current_user["plan"] != "plus":
        increment_usage(current_user["email"], "calcola")
    
    remaining = 2 - (current_user["calcola_count"] + 1) if current_user["plan"] == "pro" else "Unlimited"
    
    return {
        "renovation_cost": round(total_renovation_cost, 2),
        "total_investment": round(total_investment, 2),
        "estimated_value": round(estimated_final_value, 2),
        "roi": round(roi_percentage, 2),
        "profit": round(profit, 2),
        "duration_months": duration_months,
        "remaining_usage": remaining,
        "details": {
            "cost_sqm": round(cost_per_sqm, 2),
            "city_factor": multiplier
        }
    }

# --- NUOVO ENDPOINT PER GENERARE REPORT DOCX ---
@app.post("/features/generate-report")
async def generate_report(data: dict, current_user: dict = Depends(get_current_user)):
    # Crea un documento Word
    doc = Document()
    
    # Titolo
    title = doc.add_paragraph()
    run = title.add_run(f"Report Investimento Immobiliare - {data.get('city', 'N/A')}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6E) # Blu Big House
    
    doc.add_paragraph(f"Generato per: {current_user['name']}")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 40)
    
    # Dati Input
    doc.add_heading("Dati Immobile", level=2)
    doc.add_paragraph(f"Prezzo Acquisto: € {data.get('buy_price', 0):,.2f}")
    doc.add_paragraph(f"Superficie: {data.get('surface', 0)} mq")
    doc.add_paragraph(f"Livello Ristrutturazione: {data.get('renovation_level', 'N/A').upper()}")
    
    # Analisi Finanziaria
    doc.add_heading("Analisi Finanziaria (Stima AI)", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Costo Ristrutturazione Stimato: ").bold = True
    p.add_run(f"€ {data.get('renovation_cost', 0):,.2f}")
    
    p = doc.add_paragraph()
    p.add_run(f"Investimento Totale (inclusi oneri): ").bold = True
    p.add_run(f"€ {data.get('total_investment', 0):,.2f}")
    
    p = doc.add_paragraph()
    p.add_run(f"Valore di Rivendita Stimato: ").bold = True
    p.add_run(f"€ {data.get('estimated_value', 0):,.2f}")
    
    # Risultato
    doc.add_heading("Risultato Operazione", level=2)
    res_p = doc.add_paragraph()
    res_p.add_run(f"ROI Atteso: {data.get('roi', 0)}%").bold = True
    res_p.font.size = Pt(14)
    
    doc.add_paragraph(f"Profitto Netto Stimato: € {data.get('profit', 0):,.2f}")
    doc.add_paragraph(f"Durata Cantiere: {data.get('duration_months', 0)} mesi")
    
    # Disclaimer
    doc.add_paragraph("-" * 40)
    disclaimer = doc.add_paragraph("Disclaimer: I valori sono stime generate da AI basate su medie di mercato. Non costituiscono consulenza finanziaria.")
    disclaimer.style = "Caption"

    # Salva file temporaneo
    filename = f"report_{current_user['id']}_{int(datetime.now().timestamp())}.docx"
    doc.save(filename)
    
    return FileResponse(path=filename, filename="Report_Immobiliare_BigHouse.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.post("/features/deep-research")
async def deep_research(req: DeepResearchRequest, current_user: dict = Depends(get_current_user)):
    check_limit(current_user, "deepresearch")
    
    import time
    time.sleep(2.0) # Simula il tempo di "ragionamento" degli agenti
    
    # Simulazione risposta Agentica (CrewAI style) basata sul prompt dell'utente
    query_lower = req.query.lower()
    
    # Logica simulata per rispondere in modo contestuale
    location = "Italia"
    if "parigi" in query_lower: location = "Parigi (Francia)"
    elif "milano" in query_lower: location = "Milano"
    elif "roma" in query_lower: location = "Roma"
    
    target_roi = "standard di mercato"
    if "roi" in query_lower:
        import re
        roi_match = re.search(r'(\d+)%', req.query)
        if roi_match: target_roi = f"{roi_match.group(1)}%"

    ai_response = f"""
### 🕵️‍♂️ Report Deep Research: {req.query}

**Agente Attivo:** Senior Real Estate Analyst & Market Strategist
**Target:** {location} | **Obiettivo ROI:** {target_roi}

---

#### 1. 🌍 Analisi Macro & Location
Per raggiungere un **ROI del {target_roi}** a **{location}**, la strategia "Buy & Hold" classica non è sufficiente. È necessario orientarsi su operazioni di **Trading Immobiliare (Flipping)** o **Frazionamento**.
*   **Zone Consigliate:** { "Arrondissement 10/11 (in gentrificazione)" if "parigi" in query_lower else "Zone semi-centrali ben collegate (es. NoLo/Scalo Romana a Milano)" }.
*   **Prezzo Target Acquisto:** Necessario acquistare a sconto del 15-20% sul prezzo di mercato.

#### 2. 🏗️ Strategia di Ristrutturazione
Per massimizzare il valore:
*   **Focus:** Riqualificazione energetica (Green Homes Directive).
*   **Layout:** Trasformazione da taglio classico a layout moderno (Open space + Doppi servizi).
*   **Costo Stimato:** 1.000 - 1.200 €/mq per finiture di livello medio-alto.

#### 3. 💰 Proiezione Finanziaria (Simulata)
*   **Acquisto:** Sottostimato rispetto ai comparables.
*   **CAPEX (Lavori):** Incidenza alta ma necessaria per il repricing.
*   **Exit Strategy:** Rivendita entro 8-10 mesi.

#### 4. ⚠️ Rischi Rilevati
1.  **Tempi Burocratici:** Attenzione ai permessi per modifiche strutturali.
2.  **Volatilità Tassi:** L'aumento dei tassi potrebbe rallentare la rivendita.

---
*Analisi generata da Big House AI Agents (Market Researcher + Strategist)*
    """
    
    if current_user["plan"] != "plus":
        increment_usage(current_user["email"], "deepresearch")
    
    remaining = 2 - (current_user["deepresearch_count"] + 1) if current_user["plan"] == "pro" else "Unlimited"
    
    return {
        "result": ai_response,
        "remaining_usage": remaining
    }

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)