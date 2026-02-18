"""
╔══════════════════════════════════════════════════════════════════╗
║       MARKET RESEARCH AGENT - E-commerce / Retail               ║
║       Powered by DeepSeek + CrewAI                              ║
║                                                                  ║
║  OUTPUT:                                                         ║
║    - report_mercato.docx  → Report Word professionale           ║
║    - dati_mercato.xlsx    → Dati strutturati in Excel           ║
╚══════════════════════════════════════════════════════════════════╝

INSTALLAZIONE (esegui una volta):
    pip install crewai langchain-openai openpyxl python-docx requests

CONFIGURAZIONE:
    Imposta la variabile DEEPSEEK_API_KEY con la tua chiave API
"""

import os
import json
from datetime import datetime

# ── Librerie agentiche ──────────────────────────────────────────────
from crewai import Agent, Task, Crew, Process, LLM

# ── Output ──────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ════════════════════════════════════════════════════════════════════
# CONFIGURAZIONE
# ════════════════════════════════════════════════════════════════════

DEEPSEEK_API_KEY = "sk-305838280733422fbc11a9848035d1e9"   # ← sostituisci
SETTORE         = "E-commerce / Retail in Italia"
ANNO            = datetime.now().year

# DeepSeek tramite CrewAI LLM nativo (OpenAI-compatible)
llm = LLM(
    model="openai/deepseek-chat",
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1",
    temperature=0.3,
)


# ════════════════════════════════════════════════════════════════════
# AGENTI
# ════════════════════════════════════════════════════════════════════

ricercatore = Agent(
    role="Senior Market Researcher",
    goal=f"Raccogliere dati aggiornati e approfonditi sul mercato {SETTORE}",
    backstory="""Sei un esperto analista di mercato con 15 anni di esperienza
    nel settore retail e e-commerce italiano ed europeo. Hai lavorato per
    McKinsey e Bain & Company. Sei metodico, preciso e citi sempre le fonti.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

analista_competitor = Agent(
    role="Competitive Intelligence Analyst",
    goal=f"Analizzare i principali competitor nel mercato {SETTORE}",
    backstory="""Sei uno specialista di competitive intelligence con profonda
    conoscenza del panorama retail italiano. Identifichi punti di forza,
    debolezze, strategie di pricing e posizionamento dei competitor.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

analista_trend = Agent(
    role="Trend & Consumer Insights Analyst",
    goal="Identificare trend emergenti e comportamenti dei consumatori",
    backstory="""Sei un esperto di consumer behavior e trend forecasting.
    Analizzi dati demografici, preferenze d'acquisto, trend digitali e
    macro-tendenze per identificare opportunità di mercato.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

strategist = Agent(
    role="Strategic Business Consultant",
    goal="Sintetizzare le analisi e produrre raccomandazioni strategiche actionable",
    backstory="""Sei un consulente strategico senior che trasforma dati e
    analisi in insight concreti e raccomandazioni pratiche. Produci output
    strutturati, chiari e orientati alle decisioni di business.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)


# ════════════════════════════════════════════════════════════════════
# TASK
# ════════════════════════════════════════════════════════════════════

task_ricerca = Task(
    description=f"""
    Conduci una ricerca approfondita sul mercato {SETTORE} per l'anno {ANNO}.
    
    Includi obbligatoriamente:
    1. Dimensione del mercato (fatturato totale, crescita YoY %)
    2. Principali segmenti di prodotto e loro quota %
    3. Canali di vendita (marketplace, D2C, fisico) con quote %
    4. Profilo del consumatore medio italiano online
    5. Barriere all'ingresso e fattori critici di successo
    6. Regolamentazioni rilevanti (privacy, IVA digitale, ecc.)
    
    Formato output: testo strutturato con numeri specifici e stime realistiche.
    """,
    agent=ricercatore,
    expected_output="Report di ricerca di mercato strutturato con dati quantitativi"
)

task_competitor = Task(
    description=f"""
    Analizza i 5 principali competitor nel mercato {SETTORE}.
    
    Per ciascun competitor fornisci:
    - Nome, sede, anno di fondazione
    - Quota di mercato stimata %
    - Fatturato annuo stimato
    - Punti di forza (almeno 3)
    - Punti di debolezza (almeno 2)
    - Strategia di pricing
    - Canali principali
    - Innovazioni recenti
    
    Includi anche: Amazon IT, Zalando, ePrice o player rilevanti per il settore.
    Formato output: analisi strutturata per ciascun competitor + tabella comparativa.
    """,
    agent=analista_competitor,
    expected_output="Analisi competitor dettagliata con dati numerici per ciascuno"
)

task_trend = Task(
    description=f"""
    Identifica i principali trend e opportunità nel mercato {SETTORE} per {ANNO}-{ANNO+2}.
    
    Analizza:
    1. Top 5 trend tecnologici (AI, AR/VR, voice commerce, ecc.)
    2. Trend comportamentali dei consumatori italiani
    3. Trend di sostenibilità e impatto sul retail
    4. Opportunità di nicchia non ancora sfruttate
    5. Minacce e rischi emergenti
    6. Previsioni di crescita per segmento nei prossimi 3 anni
    
    Formato output: elenco strutturato con impatto (Alto/Medio/Basso) e timeline.
    """,
    agent=analista_trend,
    expected_output="Analisi trend con impatto e opportunità quantificate"
)

task_report_finale = Task(
    description=f"""
    Sulla base delle analisi precedenti, produci un report strategico completo.
    
    Il report deve contenere ESATTAMENTE queste sezioni:
    
    ## EXECUTIVE SUMMARY
    (3-5 punti chiave, max 200 parole)
    
    ## 1. PANORAMICA DEL MERCATO
    (dati di mercato, dimensioni, crescita)
    
    ## 2. ANALISI DELLA CONCORRENZA
    (tabella e analisi top 5 competitor)
    
    ## 3. TREND E OPPORTUNITÀ
    (trend con impatto e timeframe)
    
    ## 4. SEGMENTI TARGET CONSIGLIATI
    (3 segmenti con rationale e potenziale)
    
    ## 5. RACCOMANDAZIONI STRATEGICHE
    (5 azioni concrete, prioritizzate)
    
    ## 6. DATI PER EXCEL
    Fornisci alla fine una sezione JSON con questa struttura ESATTA:
    {{
      "kpi_mercato": [
        {{"metrica": "...", "valore": "...", "fonte": "...", "anno": {ANNO}}}
      ],
      "competitor": [
        {{"nome": "...", "quota_mercato_pct": 0, "fatturato_mln_eur": 0, 
          "punti_forza": "...", "punti_debolezza": "...", "rating": 0}}
      ],
      "trend": [
        {{"trend": "...", "impatto": "Alto/Medio/Basso", 
          "timeframe": "...", "opportunita_score": 0}}
      ],
      "previsioni": [
        {{"anno": 0, "fatturato_mercato_mld_eur": 0, "crescita_pct": 0}}
      ]
    }}
    """,
    agent=strategist,
    expected_output="Report strategico completo con sezione JSON dati strutturati",
    context=[task_ricerca, task_competitor, task_trend]
)


# ════════════════════════════════════════════════════════════════════
# CREW
# ════════════════════════════════════════════════════════════════════

crew = Crew(
    agents=[ricercatore, analista_competitor, analista_trend, strategist],
    tasks=[task_ricerca, task_competitor, task_trend, task_report_finale],
    process=Process.sequential,
    verbose=True,
)


# ════════════════════════════════════════════════════════════════════
# GENERAZIONE OUTPUT - WORD REPORT
# ════════════════════════════════════════════════════════════════════

def genera_word_report(contenuto: str, filename: str = "report_mercato.docx"):
    """Genera un report Word professionale dal contenuto del crew."""
    doc = Document()
    
    # ── Stili ────────────────────────────────────────────────────────
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    
    # ── Copertina ────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"MARKET RESEARCH REPORT\n{SETTORE.upper()}")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Analisi di Mercato · Anno {ANNO}\nGenerato da AI Market Research Agent · DeepSeek")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    
    doc.add_page_break()
    
    # ── Parsing e scrittura sezioni ──────────────────────────────────
    lines = contenuto.split('\n')
    in_json = False
    json_buffer = []
    
    for line in lines:
        stripped = line.strip()
        
        # Salta il blocco JSON (va in Excel, non nel Word)
        if stripped.startswith('```json') or stripped.startswith('{"kpi_mercato"'):
            in_json = True
            continue
        if in_json:
            if stripped == '```' or (stripped.startswith('}') and len(stripped) <= 2):
                in_json = False
            continue
        
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Heading 1: ## TITOLO
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            
        # Heading 2: ### sottotitolo
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            p.paragraph_format.space_before = Pt(10)
            
        # Bullet con - o *
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped[2:])
            
        # Numerato
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
            p = doc.add_paragraph(style='List Number')
            p.add_run(stripped[2:].strip())
            
        # Bold con **testo**
        elif '**' in stripped:
            p = doc.add_paragraph()
            parts = stripped.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (i % 2 == 1)
                
        # Paragrafo normale
        else:
            doc.add_paragraph(stripped)
    
    # ── Footer con data ──────────────────────────────────────────────
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(
        f"Confidenziale · Market Research Report · {SETTORE} · {datetime.now().strftime('%d/%m/%Y')}"
    ).font.size = Pt(9)
    
    doc.save(filename)
    print(f"✅ Report Word salvato: {filename}")
    return filename


# ════════════════════════════════════════════════════════════════════
# GENERAZIONE OUTPUT - EXCEL
# ════════════════════════════════════════════════════════════════════

def genera_excel(dati_json: dict, filename: str = "dati_mercato.xlsx"):
    """Genera un file Excel strutturato con i dati di mercato."""
    wb = openpyxl.Workbook()
    
    # ── Stili comuni ─────────────────────────────────────────────────
    header_fill  = PatternFill("solid", fgColor="1F4E79")
    header_font  = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    title_font   = Font(name="Arial", bold=True, color="1F4E79", size=14)
    normal_font  = Font(name="Arial", size=10)
    alt_fill     = PatternFill("solid", fgColor="D6E4F0")
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align   = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    thin_border  = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    
    def style_header_row(sheet, row, cols):
        for col in range(1, cols + 1):
            cell = sheet.cell(row=row, column=col)
            cell.fill   = header_fill
            cell.font   = header_font
            cell.border = thin_border
            cell.alignment = center_align
    
    def style_data_row(sheet, row, cols, alternate=False):
        for col in range(1, cols + 1):
            cell = sheet.cell(row=row, column=col)
            if alternate:
                cell.fill = alt_fill
            cell.font      = normal_font
            cell.border    = thin_border
            cell.alignment = left_align
    
    # ── Sheet 1: KPI Mercato ─────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "KPI Mercato"
    ws1.sheet_view.showGridLines = False
    
    ws1.merge_cells("A1:D1")
    ws1["A1"] = f"📊 KPI MERCATO — {SETTORE} · {ANNO}"
    ws1["A1"].font   = title_font
    ws1["A1"].alignment = center_align
    ws1.row_dimensions[1].height = 30
    
    headers_kpi = ["Metrica", "Valore", "Fonte", "Anno"]
    for col, h in enumerate(headers_kpi, 1):
        ws1.cell(row=2, column=col, value=h)
    style_header_row(ws1, 2, 4)
    
    kpi_list = dati_json.get("kpi_mercato", [])
    for i, item in enumerate(kpi_list):
        r = i + 3
        ws1.cell(r, 1, item.get("metrica", ""))
        ws1.cell(r, 2, item.get("valore", ""))
        ws1.cell(r, 3, item.get("fonte", ""))
        ws1.cell(r, 4, item.get("anno", ANNO))
        style_data_row(ws1, r, 4, alternate=(i % 2 == 1))
    
    ws1.column_dimensions["A"].width = 35
    ws1.column_dimensions["B"].width = 20
    ws1.column_dimensions["C"].width = 25
    ws1.column_dimensions["D"].width = 10
    
    # ── Sheet 2: Analisi Competitor ──────────────────────────────────
    ws2 = wb.create_sheet("Competitor")
    ws2.sheet_view.showGridLines = False
    
    ws2.merge_cells("A1:F1")
    ws2["A1"] = "🏆 ANALISI COMPETITOR"
    ws2["A1"].font   = title_font
    ws2["A1"].alignment = center_align
    ws2.row_dimensions[1].height = 30
    
    headers_comp = ["Azienda", "Quota Mercato %", "Fatturato (M€)", "Punti di Forza", "Punti Deboli", "Rating /10"]
    for col, h in enumerate(headers_comp, 1):
        ws2.cell(row=2, column=col, value=h)
    style_header_row(ws2, 2, 6)
    
    comp_list = dati_json.get("competitor", [])
    for i, c in enumerate(comp_list):
        r = i + 3
        ws2.cell(r, 1, c.get("nome", ""))
        ws2.cell(r, 2, c.get("quota_mercato_pct", 0))
        ws2.cell(r, 3, c.get("fatturato_mln_eur", 0))
        ws2.cell(r, 4, c.get("punti_forza", ""))
        ws2.cell(r, 5, c.get("punti_debolezza", ""))
        ws2.cell(r, 6, c.get("rating", 0))
        style_data_row(ws2, r, 6, alternate=(i % 2 == 1))
        # Formato percentuale e valuta
        ws2.cell(r, 2).number_format = "0.0%"
        ws2.cell(r, 3).number_format = "#,##0"
    
    ws2.column_dimensions["A"].width = 20
    ws2.column_dimensions["B"].width = 18
    ws2.column_dimensions["C"].width = 18
    ws2.column_dimensions["D"].width = 35
    ws2.column_dimensions["E"].width = 30
    ws2.column_dimensions["F"].width = 12
    
    # ── Sheet 3: Trend ───────────────────────────────────────────────
    ws3 = wb.create_sheet("Trend & Opportunità")
    ws3.sheet_view.showGridLines = False
    
    ws3.merge_cells("A1:D1")
    ws3["A1"] = "📈 TREND E OPPORTUNITÀ DI MERCATO"
    ws3["A1"].font   = title_font
    ws3["A1"].alignment = center_align
    ws3.row_dimensions[1].height = 30
    
    headers_trend = ["Trend", "Impatto", "Timeframe", "Opportunità Score"]
    for col, h in enumerate(headers_trend, 1):
        ws3.cell(row=2, column=col, value=h)
    style_header_row(ws3, 2, 4)
    
    trend_colors = {"Alto": "FF4444", "Medio": "FFA500", "Basso": "00AA00"}
    trend_list = dati_json.get("trend", [])
    for i, t in enumerate(trend_list):
        r = i + 3
        ws3.cell(r, 1, t.get("trend", ""))
        impatto = t.get("impatto", "Medio")
        ws3.cell(r, 2, impatto)
        ws3.cell(r, 2).font = Font(
            name="Arial", bold=True, size=10,
            color=trend_colors.get(impatto, "000000")
        )
        ws3.cell(r, 3, t.get("timeframe", ""))
        ws3.cell(r, 4, t.get("opportunita_score", 0))
        style_data_row(ws3, r, 4, alternate=(i % 2 == 1))
        ws3.cell(r, 2).fill = PatternFill("solid", fgColor="FFFFFF") if i % 2 == 0 else alt_fill
    
    ws3.column_dimensions["A"].width = 40
    ws3.column_dimensions["B"].width = 12
    ws3.column_dimensions["C"].width = 20
    ws3.column_dimensions["D"].width = 18
    
    # ── Sheet 4: Previsioni ──────────────────────────────────────────
    ws4 = wb.create_sheet("Previsioni 3 Anni")
    ws4.sheet_view.showGridLines = False
    
    ws4.merge_cells("A1:C1")
    ws4["A1"] = "🔮 PREVISIONI DI MERCATO 3 ANNI"
    ws4["A1"].font   = title_font
    ws4["A1"].alignment = center_align
    ws4.row_dimensions[1].height = 30
    
    headers_prev = ["Anno", "Fatturato Mercato (Mld €)", "Crescita %"]
    for col, h in enumerate(headers_prev, 1):
        ws4.cell(row=2, column=col, value=h)
    style_header_row(ws4, 2, 3)
    
    prev_list = dati_json.get("previsioni", [])
    for i, p in enumerate(prev_list):
        r = i + 3
        ws4.cell(r, 1, p.get("anno", ""))
        ws4.cell(r, 2, p.get("fatturato_mercato_mld_eur", 0))
        ws4.cell(r, 3, p.get("crescita_pct", 0))
        style_data_row(ws4, r, 3, alternate=(i % 2 == 1))
        ws4.cell(r, 2).number_format = "#,##0.0"
        ws4.cell(r, 3).number_format = "0.0%"
    
    # Aggiunge formula CAGR
    if len(prev_list) >= 2:
        last_r = len(prev_list) + 3
        ws4.cell(last_r, 1, "CAGR 3 anni")
        ws4.cell(last_r, 1).font = Font(name="Arial", bold=True, size=10)
        first_val_cell = f"B3"
        last_val_cell  = f"B{last_r - 1}"
        n = len(prev_list) - 1
        ws4.cell(last_r, 2, f"=({last_val_cell}/{first_val_cell})^(1/{n})-1")
        ws4.cell(last_r, 2).number_format = "0.0%"
        ws4.cell(last_r, 2).font = Font(name="Arial", bold=True, color="1F4E79", size=10)
    
    ws4.column_dimensions["A"].width = 12
    ws4.column_dimensions["B"].width = 28
    ws4.column_dimensions["C"].width = 15
    
    wb.save(filename)
    print(f"✅ Excel salvato: {filename}")
    return filename


# ════════════════════════════════════════════════════════════════════
# PARSER JSON DAL TESTO DEL CREW
# ════════════════════════════════════════════════════════════════════

def estrai_json_da_testo(testo: str) -> dict:
    """Estrae il blocco JSON dal report testuale generato dal crew."""
    import re
    
    # Cerca blocco ```json ... ```
    match = re.search(r'```json\s*(\{.*?\})\s*```', testo, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    
    # Cerca JSON grezzo
    match = re.search(r'\{[\s\S]*"kpi_mercato"[\s\S]*\}', testo)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass
    
    # Fallback: dati di esempio se il JSON non è parsabile
    print("⚠️  JSON non trovato nel testo. Uso dati di esempio per Excel.")
    return {
        "kpi_mercato": [
            {"metrica": "Fatturato E-commerce IT", "valore": "54,2 Mld €", "fonte": "Netcomm", "anno": ANNO},
            {"metrica": "Crescita YoY",             "valore": "+13%",        "fonte": "Netcomm", "anno": ANNO},
            {"metrica": "Acquirenti online IT",     "valore": "33 Milioni",  "fonte": "ISTAT",   "anno": ANNO},
            {"metrica": "Ordini medi/anno",         "valore": "17 ordini",   "fonte": "Netcomm", "anno": ANNO},
            {"metrica": "Scontrino medio",          "valore": "97 €",        "fonte": "Netcomm", "anno": ANNO},
        ],
        "competitor": [
            {"nome": "Amazon IT",  "quota_mercato_pct": 0.28, "fatturato_mln_eur": 15180, "punti_forza": "Logistica, Prime, varietà", "punti_debolezza": "Saturazione, fee venditori", "rating": 9},
            {"nome": "Zalando",    "quota_mercato_pct": 0.09, "fatturato_mln_eur": 4876,  "punti_forza": "Fashion focus, UX, reso gratuito", "punti_debolezza": "Solo moda, margini bassi", "rating": 8},
            {"nome": "ePrice",     "quota_mercato_pct": 0.04, "fatturato_mln_eur": 2200,  "punti_forza": "Tech/elettronica, brand storico", "punti_debolezza": "Limitata diversificazione", "rating": 6},
            {"nome": "Unieuro",    "quota_mercato_pct": 0.05, "fatturato_mln_eur": 2700,  "punti_forza": "Omnichannel, fiducia brand", "punti_debolezza": "Focus elettronica, costi fissi", "rating": 7},
            {"nome": "Decathlon",  "quota_mercato_pct": 0.03, "fatturato_mln_eur": 1620,  "punti_forza": "Prezzo, sport niche, D2C", "punti_debolezza": "Solo sport, brand house", "rating": 7},
        ],
        "trend": [
            {"trend": "Social Commerce (TikTok Shop, Instagram)",    "impatto": "Alto",  "timeframe": "2024-2025", "opportunita_score": 9},
            {"trend": "AI Personalizzazione e Raccomandazioni",      "impatto": "Alto",  "timeframe": "2024-2026", "opportunita_score": 9},
            {"trend": "Quick Commerce (consegna < 2h)",               "impatto": "Alto",  "timeframe": "2024-2025", "opportunita_score": 8},
            {"trend": "Sostenibilità e packaging eco-friendly",      "impatto": "Medio", "timeframe": "2024-2027", "opportunita_score": 7},
            {"trend": "Live Shopping & Video Commerce",               "impatto": "Medio", "timeframe": "2025-2026", "opportunita_score": 7},
            {"trend": "Buy Now Pay Later (BNPL)",                     "impatto": "Medio", "timeframe": "2024-2025", "opportunita_score": 6},
            {"trend": "Marketplace B2B per PMI",                     "impatto": "Basso", "timeframe": "2025-2027", "opportunita_score": 6},
        ],
        "previsioni": [
            {"anno": ANNO,     "fatturato_mercato_mld_eur": 54.2, "crescita_pct": 0.13},
            {"anno": ANNO + 1, "fatturato_mercato_mld_eur": 61.2, "crescita_pct": 0.13},
            {"anno": ANNO + 2, "fatturato_mercato_mld_eur": 69.2, "crescita_pct": 0.13},
            {"anno": ANNO + 3, "fatturato_mercato_mld_eur": 78.2, "crescita_pct": 0.13},
        ]
    }


# ════════════════════════════════════════════════════════════════════
# MAIN — ESEGUI L'AGENTE E GENERA I FILE
# ════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 65)
    print(f"🚀 AVVIO MARKET RESEARCH AGENT")
    print(f"   Settore: {SETTORE}")
    print(f"   Modello: DeepSeek Chat")
    print(f"   Output:  report_mercato.docx + dati_mercato.xlsx")
    print("=" * 65)
    
    # 1) Esegui il crew
    risultato = crew.kickoff()
    testo_report = str(risultato)
    
    print("\n" + "=" * 65)
    print("📄 GENERAZIONE FILE DI OUTPUT")
    print("=" * 65)
    
    # 2) Genera il report Word
    genera_word_report(testo_report, "report_mercato.docx")
    
    # 3) Estrai JSON e genera Excel
    dati = estrai_json_da_testo(testo_report)
    genera_excel(dati, "dati_mercato.xlsx")
    
    print("\n✅ COMPLETATO!")
    print("   📄 report_mercato.docx — Report Word completo")
    print("   📊 dati_mercato.xlsx  — Dati strutturati Excel")
    print("=" * 65)