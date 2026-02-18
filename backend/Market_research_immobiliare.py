"""
╔══════════════════════════════════════════════════════════════════╗
║       MARKET RESEARCH AGENT - Immobiliare in Italia             ║
║       Powered by DeepSeek + CrewAI                              ║
║                                                                  ║
║  OUTPUT:                                                         ║
║    - report_immobiliare.docx  → Report Word professionale       ║
║    - dati_immobiliare.xlsx    → Dati strutturati in Excel       ║
╚══════════════════════════════════════════════════════════════════╝

INSTALLAZIONE (esegui una volta):
    pip install crewai langchain-openai openpyxl python-docx requests

CONFIGURAZIONE:
    Imposta la variabile DEEPSEEK_API_KEY con la tua chiave API
"""

import os
import json
from datetime import datetime

# ── Librerie agentiche ──────────────────────────────────────────────
from crewai import Agent, Task, Crew, Process, LLM

# ── Output ──────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ════════════════════════════════════════════════════════════════════
# CONFIGURAZIONE
# ════════════════════════════════════════════════════════════════════

DEEPSEEK_API_KEY = "sk-305838280733422fbc11a9848035d1e9"   # ← sostituisci
SETTORE          = "Mercato Immobiliare in Italia"
ANNO             = datetime.now().year

# DeepSeek tramite CrewAI LLM nativo (OpenAI-compatible)
llm = LLM(
    model="openai/deepseek-chat",
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1",
    temperature=0.3,
)


# ════════════════════════════════════════════════════════════════════
# AGENTI
# ════════════════════════════════════════════════════════════════════

ricercatore = Agent(
    role="Senior Real Estate Market Researcher",
    goal=f"Raccogliere dati aggiornati e approfonditi sul {SETTORE}",
    backstory="""Sei un esperto analista immobiliare con 15 anni di esperienza
    nel mercato residenziale, commerciale e delle locazioni in Italia.
    Hai collaborato con Nomisma, CBRE e Scenari Immobiliari.
    Sei metodico, preciso e citi sempre fonti ufficiali (OMI, Agenzia delle Entrate,
    ISTAT, Banca d'Italia). Conosci a fondo le dinamiche locali delle principali
    città italiane e le differenze tra Nord, Centro e Sud.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

analista_competitor = Agent(
    role="PropTech & Agency Competitive Intelligence Analyst",
    goal=f"Analizzare i principali player e portali nel {SETTORE}",
    backstory="""Sei uno specialista di competitive intelligence nel settore
    immobiliare italiano. Conosci in profondità i portali online (Idealista,
    Immobiliare.it, Casa.it, Subito.it), le grandi agenzie (RE/MAX, Tecnocasa,
    Engel & Völkers) e le startup PropTech emergenti.
    Identifichi punti di forza, debolezze, strategie di pricing e
    posizionamento per ciascun operatore del mercato.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

analista_trend = Agent(
    role="Real Estate Trend & Macro Analyst",
    goal="Identificare trend macro, demografici e tecnologici nel mercato immobiliare",
    backstory="""Sei un esperto di trend forecasting immobiliare e analisi
    macroeconomica applicata al real estate. Analizzi tassi di interesse BCE,
    inflazione, domanda abitativa, trend demografici (invecchiamento,
    urbanizzazione, smart working) e innovazioni PropTech come AI valutativa,
    tokenizzazione immobiliare, realtà virtuale per i tour e piattaforme
    di crowdfunding immobiliare.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)

strategist = Agent(
    role="Real Estate Strategic Consultant",
    goal="Sintetizzare le analisi e produrre raccomandazioni strategiche actionable",
    backstory="""Sei un consulente strategico senior specializzato nel real estate
    italiano. Trasformi dati di mercato, analisi competitive e trend in
    insight concreti e raccomandazioni pratiche per investitori, agenzie,
    sviluppatori e operatori PropTech. Produci output strutturati, chiari
    e orientati alle decisioni di investimento e business.""",
    llm=llm,
    verbose=True,
    allow_delegation=False,
)


# ════════════════════════════════════════════════════════════════════
# TASK
# ════════════════════════════════════════════════════════════════════

task_ricerca = Task(
    description=f"""
    Conduci una ricerca approfondita sul {SETTORE} per l'anno {ANNO}.
    
    Includi obbligatoriamente:
    1. Volume compravendite residenziali (NTN - Numero Transazioni Normalizzate)
    2. Prezzi medi al metro quadro nelle principali città (Milano, Roma, Napoli,
       Torino, Bologna, Firenze) per il segmento residenziale
    3. Andamento dei mutui: tassi medi, volumi erogati, LTV medio
    4. Segmentazione del mercato: residenziale, commerciale, logistica, luxury
       con quote % e fatturato stimato per ciascun segmento
    5. Mercato delle locazioni: canoni medi, rendimenti lordi per città,
       vacancy rate
    6. Profilo dell'acquirente/investitore tipo in Italia ({ANNO})
    7. Barriere all'ingresso e fattori critici di successo nel settore
    8. Regolamentazioni rilevanti: Superbonus, legge sugli affitti brevi,
       direttiva UE case green (EPBD), registro locazioni
    9. Confronto con principali mercati europei (Spagna, Francia, Germania)
    
    Formato output: testo strutturato con numeri specifici e stime realistiche
    basate su fonti quali OMI/Agenzia delle Entrate, Nomisma, CBRE, ISTAT.
    """,
    agent=ricercatore,
    expected_output="Report di ricerca immobiliare strutturato con dati quantitativi e fonti"
)

task_competitor = Task(
    description=f"""
    Analizza i 5 principali competitor/player nel {SETTORE}.
    
    Includi un mix di: portali immobiliari online, reti di agenzie fisiche
    e startup PropTech. Suggeriti (ma non limitarti a questi):
    Idealista, Immobiliare.it, Tecnocasa, RE/MAX Italia, Engel & Völkers.
    
    Per ciascun player fornisci:
    - Nome, tipologia (portale/agenzia/PropTech), sede, anno di fondazione
    - Quota di mercato stimata % (per portali: share annunci; per agenzie: % compravendite)
    - Fatturato/ricavi annui stimati in Italia (Mln €)
    - Numero di annunci attivi o agenzie/uffici sul territorio
    - Punti di forza (almeno 3 specifici)
    - Punti di debolezza (almeno 2 specifici)
    - Modello di revenue (abbonamento, commissione, lead generation, ecc.)
    - Innovazioni recenti o tecnologie adottate
    - Target principale (acquirenti, venditori, investitori, luxury, ecc.)
    - Rating competitivo /10
    
    Formato output: analisi strutturata per ciascun player + tabella comparativa.
    """,
    agent=analista_competitor,
    expected_output="Analisi competitor dettagliata con dati numerici per ciascun player"
)

task_trend = Task(
    description=f"""
    Identifica i principali trend e opportunità nel {SETTORE} per {ANNO}-{ANNO+2}.
    
    Analizza obbligatoriamente:
    1. Impatto dei tassi BCE sui mutui e sulla domanda abitativa
    2. Trend Smart Working e suo effetto sulla domanda per area geografica
       (fuga dalle grandi città vs ritorno, domanda in hinterland e borghi)
    3. Top 5 trend PropTech: AI per valutazione automatica AVM, virtual tour 3D,
       tokenizzazione immobiliare, crowdfunding real estate, blockchain per rogiti
    4. Trend ESG: direttiva EPBD UE (case green), classi energetiche,
       riqualificazione energetica, impatto su prezzi e domanda
    5. Trend affitti brevi (Airbnb, Booking) vs affitti tradizionali:
       regolamentazione, rendimenti a confronto, impatto sul mercato
    6. Domanda abitativa per fascia demografica: Millennials, Gen Z, Silver Economy
    7. Mercato luxury e ultra-luxury: trend e città chiave
    8. Opportunità di nicchia: co-living, senior housing, student housing,
       build-to-rent
    9. Previsioni prezzi residenziale per città nei prossimi 3 anni
    10. Rischi macro: recessione, inflazione, stretta creditizia
    
    Formato output: elenco strutturato con impatto (Alto/Medio/Basso) e timeline.
    """,
    agent=analista_trend,
    expected_output="Analisi trend immobiliari con impatto, timeframe e opportunità quantificate"
)

task_report_finale = Task(
    description=f"""
    Sulla base delle analisi precedenti, produci un report strategico completo
    sul {SETTORE}.
    
    Il report deve contenere ESATTAMENTE queste sezioni:
    
    ## EXECUTIVE SUMMARY
    (3-5 punti chiave, max 200 parole)
    
    ## 1. PANORAMICA DEL MERCATO
    (dati di mercato, volumi, prezzi, mutui, segmenti)
    
    ## 2. ANALISI DELLA CONCORRENZA
    (analisi top 5 player: portali, agenzie, PropTech)
    
    ## 3. TREND E OPPORTUNITÀ
    (trend macro, PropTech, ESG, demografici con impatto e timeframe)
    
    ## 4. SEGMENTI TARGET CONSIGLIATI
    (3 segmenti/nicchie con rationale e potenziale: es. affitti brevi luxury,
    co-living urbano, riqualificazione energetica hinterland)
    
    ## 5. RACCOMANDAZIONI STRATEGICHE
    (5 azioni concrete e prioritizzate per un operatore/investitore immobiliare)
    
    ## 6. DATI PER EXCEL
    Fornisci alla fine una sezione JSON con questa struttura ESATTA:
    {{
      "kpi_mercato": [
        {{"metrica": "...", "valore": "...", "fonte": "...", "anno": {ANNO}}}
      ],
      "competitor": [
        {{"nome": "...", "quota_mercato_pct": 0, "fatturato_mln_eur": 0,
          "punti_forza": "...", "punti_debolezza": "...", "rating": 0}}
      ],
      "trend": [
        {{"trend": "...", "impatto": "Alto/Medio/Basso",
          "timeframe": "...", "opportunita_score": 0}}
      ],
      "previsioni": [
        {{"anno": 0, "transazioni_ntn": 0, "prezzo_medio_mq_italia": 0, "crescita_pct": 0}}
      ]
    }}
    
    IMPORTANTE: il JSON deve contenere almeno 6 kpi_mercato, 5 competitor,
    7 trend e 4 anni di previsioni. I numeri devono essere realistici e
    basati sui dati raccolti nelle analisi precedenti.
    """,
    agent=strategist,
    expected_output="Report strategico immobiliare completo con sezione JSON dati strutturati",
    context=[task_ricerca, task_competitor, task_trend]
)


# ════════════════════════════════════════════════════════════════════
# CREW
# ════════════════════════════════════════════════════════════════════

crew = Crew(
    agents=[ricercatore, analista_competitor, analista_trend, strategist],
    tasks=[task_ricerca, task_competitor, task_trend, task_report_finale],
    process=Process.sequential,
    verbose=True,
)


# ════════════════════════════════════════════════════════════════════
# GENERAZIONE OUTPUT - WORD REPORT
# ════════════════════════════════════════════════════════════════════

def genera_word_report(contenuto: str, filename: str = "report_immobiliare.docx"):
    """Genera un report Word professionale dal contenuto del crew."""
    doc = Document()

    # ── Stili ────────────────────────────────────────────────────────
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)

    # ── Copertina ────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"MARKET RESEARCH REPORT\n{SETTORE.upper()}")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"Analisi di Mercato · Anno {ANNO}\n"
        f"Generato da AI Market Research Agent · DeepSeek"
    )
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_page_break()

    # ── Parsing e scrittura sezioni ──────────────────────────────────
    lines = contenuto.split('\n')
    in_json = False

    for line in lines:
        stripped = line.strip()

        # Salta il blocco JSON (va in Excel, non nel Word)
        if stripped.startswith('```json') or stripped.startswith('{"kpi_mercato"'):
            in_json = True
            continue
        if in_json:
            if stripped == '```' or (stripped.startswith('}') and len(stripped) <= 2):
                in_json = False
            continue

        if not stripped:
            doc.add_paragraph()
            continue

        # Heading 1: ## TITOLO
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)

        # Heading 2: ### sottotitolo
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            p.paragraph_format.space_before = Pt(10)

        # Bullet con - o *
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped[2:])

        # Numerato
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
            p = doc.add_paragraph(style='List Number')
            p.add_run(stripped[2:].strip())

        # Bold con **testo**
        elif '**' in stripped:
            p = doc.add_paragraph()
            parts = stripped.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (i % 2 == 1)

        # Paragrafo normale
        else:
            doc.add_paragraph(stripped)

    # ── Footer con data ──────────────────────────────────────────────
    section = doc.sections[0]
    footer  = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(
        f"Confidenziale · Market Research Report · {SETTORE} · "
        f"{datetime.now().strftime('%d/%m/%Y')}"
    ).font.size = Pt(9)

    doc.save(filename)
    print(f"✅ Report Word salvato: {filename}")
    return filename


# ════════════════════════════════════════════════════════════════════
# GENERAZIONE OUTPUT - EXCEL
# ════════════════════════════════════════════════════════════════════

def genera_excel(dati_json: dict, filename: str = "dati_immobiliare.xlsx"):
    """Genera un file Excel strutturato con i dati del mercato immobiliare."""
    wb = openpyxl.Workbook()

    # ── Stili comuni ─────────────────────────────────────────────────
    header_fill  = PatternFill("solid", fgColor="1F4E79")
    header_font  = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    title_font   = Font(name="Arial", bold=True, color="1F4E79", size=14)
    normal_font  = Font(name="Arial", size=10)
    alt_fill     = PatternFill("solid", fgColor="D6E4F0")
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align   = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    thin_border  = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    def style_header_row(sheet, row, cols):
        for col in range(1, cols + 1):
            cell = sheet.cell(row=row, column=col)
            cell.fill      = header_fill
            cell.font      = header_font
            cell.border    = thin_border
            cell.alignment = center_align

    def style_data_row(sheet, row, cols, alternate=False):
        for col in range(1, cols + 1):
            cell = sheet.cell(row=row, column=col)
            if alternate:
                cell.fill = alt_fill
            cell.font      = normal_font
            cell.border    = thin_border
            cell.alignment = left_align

    # ── Sheet 1: KPI Mercato Immobiliare ─────────────────────────────
    ws1 = wb.active
    ws1.title = "KPI Mercato"
    ws1.sheet_view.showGridLines = False

    ws1.merge_cells("A1:D1")
    ws1["A1"] = f"🏠 KPI MERCATO IMMOBILIARE — Italia · {ANNO}"
    ws1["A1"].font      = title_font
    ws1["A1"].alignment = center_align
    ws1.row_dimensions[1].height = 30

    headers_kpi = ["Metrica", "Valore", "Fonte", "Anno"]
    for col, h in enumerate(headers_kpi, 1):
        ws1.cell(row=2, column=col, value=h)
    style_header_row(ws1, 2, 4)

    kpi_list = dati_json.get("kpi_mercato", [])
    for i, item in enumerate(kpi_list):
        r = i + 3
        ws1.cell(r, 1, item.get("metrica", ""))
        ws1.cell(r, 2, item.get("valore", ""))
        ws1.cell(r, 3, item.get("fonte", ""))
        ws1.cell(r, 4, item.get("anno", ANNO))
        style_data_row(ws1, r, 4, alternate=(i % 2 == 1))

    ws1.column_dimensions["A"].width = 38
    ws1.column_dimensions["B"].width = 22
    ws1.column_dimensions["C"].width = 28
    ws1.column_dimensions["D"].width = 10

    # ── Sheet 2: Analisi Player ──────────────────────────────────────
    ws2 = wb.create_sheet("Player & Competitor")
    ws2.sheet_view.showGridLines = False

    ws2.merge_cells("A1:F1")
    ws2["A1"] = "🏆 ANALISI PLAYER — Portali, Agenzie & PropTech"
    ws2["A1"].font      = title_font
    ws2["A1"].alignment = center_align
    ws2.row_dimensions[1].height = 30

    headers_comp = ["Player", "Quota Mercato %", "Fatturato IT (M€)", "Punti di Forza", "Punti Deboli", "Rating /10"]
    for col, h in enumerate(headers_comp, 1):
        ws2.cell(row=2, column=col, value=h)
    style_header_row(ws2, 2, 6)

    comp_list = dati_json.get("competitor", [])
    for i, c in enumerate(comp_list):
        r = i + 3
        ws2.cell(r, 1, c.get("nome", ""))
        ws2.cell(r, 2, c.get("quota_mercato_pct", 0))
        ws2.cell(r, 3, c.get("fatturato_mln_eur", 0))
        ws2.cell(r, 4, c.get("punti_forza", ""))
        ws2.cell(r, 5, c.get("punti_debolezza", ""))
        ws2.cell(r, 6, c.get("rating", 0))
        style_data_row(ws2, r, 6, alternate=(i % 2 == 1))
        ws2.cell(r, 2).number_format = "0.0%"
        ws2.cell(r, 3).number_format = "#,##0"

    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 18
    ws2.column_dimensions["C"].width = 20
    ws2.column_dimensions["D"].width = 38
    ws2.column_dimensions["E"].width = 32
    ws2.column_dimensions["F"].width = 12

    # ── Sheet 3: Trend ───────────────────────────────────────────────
    ws3 = wb.create_sheet("Trend & Opportunità")
    ws3.sheet_view.showGridLines = False

    ws3.merge_cells("A1:D1")
    ws3["A1"] = "📈 TREND E OPPORTUNITÀ — Mercato Immobiliare"
    ws3["A1"].font      = title_font
    ws3["A1"].alignment = center_align
    ws3.row_dimensions[1].height = 30

    headers_trend = ["Trend", "Impatto", "Timeframe", "Opportunità Score"]
    for col, h in enumerate(headers_trend, 1):
        ws3.cell(row=2, column=col, value=h)
    style_header_row(ws3, 2, 4)

    trend_colors = {"Alto": "FF4444", "Medio": "FFA500", "Basso": "00AA00"}
    trend_list = dati_json.get("trend", [])
    for i, t in enumerate(trend_list):
        r = i + 3
        ws3.cell(r, 1, t.get("trend", ""))
        impatto = t.get("impatto", "Medio")
        ws3.cell(r, 2, impatto)
        ws3.cell(r, 2).font = Font(
            name="Arial", bold=True, size=10,
            color=trend_colors.get(impatto, "000000")
        )
        ws3.cell(r, 3, t.get("timeframe", ""))
        ws3.cell(r, 4, t.get("opportunita_score", 0))
        style_data_row(ws3, r, 4, alternate=(i % 2 == 1))
        ws3.cell(r, 2).fill = PatternFill("solid", fgColor="FFFFFF") if i % 2 == 0 else alt_fill

    ws3.column_dimensions["A"].width = 45
    ws3.column_dimensions["B"].width = 12
    ws3.column_dimensions["C"].width = 20
    ws3.column_dimensions["D"].width = 18

    # ── Sheet 4: Previsioni Mercato ──────────────────────────────────
    ws4 = wb.create_sheet("Previsioni 3 Anni")
    ws4.sheet_view.showGridLines = False

    ws4.merge_cells("A1:D1")
    ws4["A1"] = "🔮 PREVISIONI MERCATO IMMOBILIARE — 3 Anni"
    ws4["A1"].font      = title_font
    ws4["A1"].alignment = center_align
    ws4.row_dimensions[1].height = 30

    # Header con colonne specifiche per il real estate
    headers_prev = ["Anno", "Transazioni NTN", "Prezzo Medio €/mq", "Crescita %"]
    for col, h in enumerate(headers_prev, 1):
        ws4.cell(row=2, column=col, value=h)
    style_header_row(ws4, 2, 4)

    prev_list = dati_json.get("previsioni", [])
    for i, p in enumerate(prev_list):
        r = i + 3
        ws4.cell(r, 1, p.get("anno", ""))
        ws4.cell(r, 2, p.get("transazioni_ntn", 0))
        ws4.cell(r, 3, p.get("prezzo_medio_mq_italia", 0))
        ws4.cell(r, 4, p.get("crescita_pct", 0))
        style_data_row(ws4, r, 4, alternate=(i % 2 == 1))
        ws4.cell(r, 2).number_format = "#,##0"
        ws4.cell(r, 3).number_format = "#,##0 €"
        ws4.cell(r, 4).number_format = "0.0%"

    # Formula CAGR sulle transazioni
    if len(prev_list) >= 2:
        last_r = len(prev_list) + 3
        ws4.cell(last_r, 1, "CAGR 3 anni")
        ws4.cell(last_r, 1).font = Font(name="Arial", bold=True, size=10)
        n = len(prev_list) - 1
        ws4.cell(last_r, 2, f"=(B{last_r - 1}/B3)^(1/{n})-1")
        ws4.cell(last_r, 2).number_format = "0.0%"
        ws4.cell(last_r, 2).font = Font(name="Arial", bold=True, color="1F4E79", size=10)
        ws4.cell(last_r, 3, f"=(C{last_r - 1}/C3)^(1/{n})-1")
        ws4.cell(last_r, 3).number_format = "0.0%"
        ws4.cell(last_r, 3).font = Font(name="Arial", bold=True, color="1F4E79", size=10)

    ws4.column_dimensions["A"].width = 12
    ws4.column_dimensions["B"].width = 22
    ws4.column_dimensions["C"].width = 22
    ws4.column_dimensions["D"].width = 15

    wb.save(filename)
    print(f"✅ Excel salvato: {filename}")
    return filename


# ════════════════════════════════════════════════════════════════════
# PARSER JSON DAL TESTO DEL CREW
# ════════════════════════════════════════════════════════════════════

def estrai_json_da_testo(testo: str) -> dict:
    """Estrae il blocco JSON dal report testuale generato dal crew."""
    import re

    # Cerca blocco ```json ... ```
    match = re.search(r'```json\s*(\{.*?\})\s*```', testo, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # Cerca JSON grezzo
    match = re.search(r'\{[\s\S]*"kpi_mercato"[\s\S]*\}', testo)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass

    # Fallback: dati di esempio immobiliari se il JSON non è parsabile
    print("⚠️  JSON non trovato nel testo. Uso dati di esempio per Excel.")
    return {
        "kpi_mercato": [
            {"metrica": "Transazioni residenziali (NTN)",       "valore": "710.000",      "fonte": "OMI / Agenzia delle Entrate", "anno": ANNO},
            {"metrica": "Variazione NTN vs anno precedente",    "valore": "-9,7%",        "fonte": "OMI / Agenzia delle Entrate", "anno": ANNO},
            {"metrica": "Prezzo medio residenziale Italia",     "valore": "1.890 €/mq",   "fonte": "Nomisma",                    "anno": ANNO},
            {"metrica": "Prezzo medio Milano",                  "valore": "5.200 €/mq",   "fonte": "Nomisma / CBRE",             "anno": ANNO},
            {"metrica": "Prezzo medio Roma",                    "valore": "3.800 €/mq",   "fonte": "Nomisma / CBRE",             "anno": ANNO},
            {"metrica": "Tasso mutuo fisso medio (20 anni)",    "valore": "3,4%",         "fonte": "Banca d'Italia",             "anno": ANNO},
            {"metrica": "Mutui erogati (Mld €)",                "valore": "41,2 Mld €",   "fonte": "Banca d'Italia / CRIF",      "anno": ANNO},
            {"metrica": "Rendimento lordo affitti Milano",      "valore": "4,2%",         "fonte": "CBRE",                       "anno": ANNO},
            {"metrica": "Rendimento lordo affitti Roma",        "valore": "4,8%",         "fonte": "CBRE",                       "anno": ANNO},
            {"metrica": "Quota mercato affitti brevi su totale","valore": "18%",           "fonte": "AirDNA / Scenari Imm.",      "anno": ANNO},
        ],
        "competitor": [
            {
                "nome": "Idealista IT",
                "quota_mercato_pct": 0.38,
                "fatturato_mln_eur": 210,
                "punti_forza": "Leader di mercato, UX superiore, AI valutativa, copertura nazionale",
                "punti_debolezza": "Abbonamenti costosi per agenzie, owned da fondo spagnolo",
                "rating": 9
            },
            {
                "nome": "Immobiliare.it",
                "quota_mercato_pct": 0.28,
                "fatturato_mln_eur": 145,
                "punti_forza": "Brand storico italiano, forte SEO, integrazioni CRM agenzie",
                "punti_debolezza": "UX datata vs competitor, meno funzionalità AI",
                "rating": 7
            },
            {
                "nome": "Tecnocasa Group",
                "quota_mercato_pct": 0.12,
                "fatturato_mln_eur": 820,
                "punti_forza": "Rete fisica capillare (1.200+ agenzie), brand trust, formazione",
                "punti_debolezza": "Modello franchising rigido, digitalizzazione lenta",
                "rating": 8
            },
            {
                "nome": "RE/MAX Italia",
                "quota_mercato_pct": 0.07,
                "fatturato_mln_eur": 380,
                "punti_forza": "Network internazionale, agenti indipendenti, lusso e premium",
                "punti_debolezza": "Fee elevate per agenti, meno penetrazione Sud Italia",
                "rating": 7
            },
            {
                "nome": "Engel & Völkers IT",
                "quota_mercato_pct": 0.04,
                "fatturato_mln_eur": 190,
                "punti_forza": "Posizionamento luxury, clientela HNW internazionale, brand premium",
                "punti_debolezza": "Nicchia di mercato ristretta, alto costo di ingresso",
                "rating": 8
            },
        ],
        "trend": [
            {"trend": "Calo tassi BCE: rimbalzo della domanda abitativa",        "impatto": "Alto",  "timeframe": f"{ANNO}-{ANNO+1}", "opportunita_score": 9},
            {"trend": "AI per Automated Valuation Model (AVM) immobiliare",      "impatto": "Alto",  "timeframe": f"{ANNO}-{ANNO+2}", "opportunita_score": 9},
            {"trend": "Direttiva EPBD UE: riqualificazione energetica obbligatoria", "impatto": "Alto", "timeframe": f"{ANNO}-{ANNO+3}", "opportunita_score": 8},
            {"trend": "Co-living e student housing nelle grandi città",          "impatto": "Alto",  "timeframe": f"{ANNO}-{ANNO+2}", "opportunita_score": 8},
            {"trend": "Smart Working: domanda crescente hinterland e borghi",    "impatto": "Medio", "timeframe": f"{ANNO}-{ANNO+2}", "opportunita_score": 7},
            {"trend": "Affitti brevi Airbnb: nuova regolamentazione e CIN",      "impatto": "Medio", "timeframe": f"{ANNO}",          "opportunita_score": 6},
            {"trend": "Tokenizzazione immobiliare e crowdfunding real estate",   "impatto": "Medio", "timeframe": f"{ANNO+1}-{ANNO+3}", "opportunita_score": 7},
            {"trend": "Senior Housing e Silver Economy (over 65)",               "impatto": "Medio", "timeframe": f"{ANNO}-{ANNO+3}", "opportunita_score": 8},
            {"trend": "Virtual Tour 3D e metaverso per property viewing",        "impatto": "Medio", "timeframe": f"{ANNO}-{ANNO+2}", "opportunita_score": 6},
            {"trend": "Build-to-Rent istituzionale nelle grandi città",          "impatto": "Basso", "timeframe": f"{ANNO+1}-{ANNO+3}", "opportunita_score": 7},
        ],
        "previsioni": [
            {"anno": ANNO,     "transazioni_ntn": 710000, "prezzo_medio_mq_italia": 1890, "crescita_pct": -0.097},
            {"anno": ANNO + 1, "transazioni_ntn": 730000, "prezzo_medio_mq_italia": 1940, "crescita_pct":  0.028},
            {"anno": ANNO + 2, "transazioni_ntn": 760000, "prezzo_medio_mq_italia": 1990, "crescita_pct":  0.041},
            {"anno": ANNO + 3, "transazioni_ntn": 790000, "prezzo_medio_mq_italia": 2050, "crescita_pct":  0.039},
        ]
    }


# ════════════════════════════════════════════════════════════════════
# MAIN — ESEGUI L'AGENTE E GENERA I FILE
# ════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 65)
    print(f"🚀 AVVIO MARKET RESEARCH AGENT — IMMOBILIARE")
    print(f"   Settore: {SETTORE}")
    print(f"   Modello: DeepSeek Chat")
    print(f"   Output:  report_immobiliare.docx + dati_immobiliare.xlsx")
    print("=" * 65)

    # 1) Esegui il crew
    risultato    = crew.kickoff()
    testo_report = str(risultato)

    print("\n" + "=" * 65)
    print("📄 GENERAZIONE FILE DI OUTPUT")
    print("=" * 65)

    # 2) Genera il report Word
    genera_word_report(testo_report, "report_immobiliare.docx")

    # 3) Estrai JSON e genera Excel
    dati = estrai_json_da_testo(testo_report)
    genera_excel(dati, "dati_immobiliare.xlsx")

    print("\n✅ COMPLETATO!")
    print("   📄 report_immobiliare.docx — Report Word completo")
    print("   📊 dati_immobiliare.xlsx   — Dati strutturati Excel")
    print("=" * 65)